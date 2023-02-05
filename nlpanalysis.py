import os
import openai
import json
from docx import Document



## These are the prompts that will be asked through the API

# MARKDOWN_TABLE_PROMPT = "Create a table of Key Accomplishment, Steps to Achievement, Evaluation/Measurement, Key Indicators, Key Learnings from the following text: "

MAIN_TOPIC_PROMP = "What is the main topic of this document?*:"
KEY_ACCOMPLISHMENTS_PROMPT = "What are the key accomplishments of the following text has?*:"
CHALLENGES_PROMPT = "What are the challenges of the following text has?*:"
KEY_LEARNINGS_PROMPT = "What are the key learnings of the following text has?*:"
STATISTICS_PROMPT = "What are the main statistics of the following text has?*:"
SUMMARY_PROMPT = "What is the summary of this document?*:"
SOMETHING_INTERESTING = "What is something you found insightful of the following text?*:"


## This is the API key that will be used to access the OpenAI API
openai.api_key = "YOUR_KEY"

## This is the raw input data that will be used as an input data for the prompts above. 
WHITE_SPACE = "                                     "


##  open a text fild and store it on a variable called data_string
document = Document()

document.add_heading('DEISHACK: ACCESSIBLE NLP TOOL', 0)

file = open("data2.txt", "r")
data_string = " "

document.add_paragraph("This is the output of the document from the langauge model.")
try:
  for line in file:
    data_string = data_string + line + " "
except UnicodeDecodeError:
  pass

### Creation of Document ###
file = open("output.txt", 'w')

file.write("This is the output of the document from the langauge model.")
file.write(WHITE_SPACE)
file.write(WHITE_SPACE)
file.write(WHITE_SPACE)

####  MODELS BELOW  #### 

RAW_INPUT_DATA = data_string



## This function will generate a topic sentence for the document
def ask_davinci(Full_Prompt):
  response = openai.Completion.create(
    model="text-davinci-003",
    prompt=Full_Prompt,
    temperature=0.7,
    max_tokens=1006,
    top_p=1,
    frequency_penalty=0,
    presence_penalty=0
  )
  return response 


### GENERATING MAIN TOPIC

INPUT_Generate_topic_sentence = MAIN_TOPIC_PROMP + WHITE_SPACE + RAW_INPUT_DATA
response = ask_davinci(INPUT_Generate_topic_sentence)
RESPONSE_text_topic_sentence = response["choices"][0]["text"]
Heading_main_topic = "The Main Topic of the Document is:"

print("<<<<< THE MAIN TOPIC OF THE INPUT IS: >>>>>")
print("")
print(RESPONSE_text_topic_sentence)
print("")


document.add_heading(Heading_main_topic, level=1)
document.add_paragraph(RESPONSE_text_topic_sentence)


### GENERATING KEY ACCOMPLISHMENTS
INPUT_Markdown_Table = KEY_ACCOMPLISHMENTS_PROMPT + WHITE_SPACE + RAW_INPUT_DATA
response = ask_davinci(INPUT_Markdown_Table)
RESPONSE_key_accomplishments = response["choices"][0]["text"]
Heading_key_accomplishments = "The Key Accomplishments of the Document are:"

print("<<<<< THE KEY ACCOMPLISHMENTS ARE: >>>>>")
print("")
print(RESPONSE_key_accomplishments)
print("")
document.add_heading(Heading_key_accomplishments, level=1)
document.add_paragraph(RESPONSE_key_accomplishments)


### GENERATING CHALLENGES 
INPUT_Summary = CHALLENGES_PROMPT + WHITE_SPACE + RAW_INPUT_DATA
response = ask_davinci(INPUT_Summary)
RESPONSE_text_summary = response["choices"][0]["text"]
Heading_challenges = "The Challenges of the Document are:"

print("<<<<< THE MAIN CHALLENGES: >>>>>")
print("")
print(RESPONSE_text_summary)
print("")
document.add_heading(Heading_challenges, level=1)
document.add_paragraph(RESPONSE_text_summary)


### GENERATING STATISTICS
INPUT_Main_Statistics = STATISTICS_PROMPT + WHITE_SPACE + RAW_INPUT_DATA
response = ask_davinci(INPUT_Main_Statistics)
RESPONSE_text_statistics = response["choices"][0]["text"]
Heading_statistics = "The Main Statistics of the Document are:"

print("<<<<< THESE ARE THE MAIN STATISTICS: >>>>>")
print("")
print(RESPONSE_text_statistics)
print("")
document.add_heading(Heading_statistics, level=1)
document.add_paragraph(RESPONSE_text_statistics)


### GENERATING FULL SUMMARY
INPUT_Summary = SUMMARY_PROMPT + WHITE_SPACE + RAW_INPUT_DATA
response = ask_davinci(INPUT_Summary)
RESPONSE_text_summary = response["choices"][0]["text"]
Heading_summary = "The Summary of the Document is:"

print("<<<<< FULL SUMMARY: >>>>>")
print("")
print(RESPONSE_text_summary)
print("")
document.add_heading(Heading_summary, level=1)
document.add_paragraph(RESPONSE_text_summary)

#### Generatin something interesting
INPUT_Summary = SOMETHING_INTERESTING + WHITE_SPACE + RAW_INPUT_DATA
response = ask_davinci(INPUT_Summary)
RESPONSE_interesting = response["choices"][0]["text"]
Heading_interesting = "Something intersting the model found from this paper is:"

print("<<<<< Something Interesting: >>>>>")
print("")
print(RESPONSE_interesting)
print("")
document.add_heading(Heading_interesting, level=1)
document.add_paragraph(RESPONSE_interesting)



## Saving Output Document
document.save('demo.docx')





